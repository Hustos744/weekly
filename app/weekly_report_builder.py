import json
import os
import re
import tempfile
import threading
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# -------------------------
# Config loading/saving
# -------------------------

DEFAULT_CONFIG_PATH = os.path.join(os.getcwd(), "config.json")


def load_config(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    validate_config(cfg)
    return cfg


def save_config(path: str, cfg: dict) -> None:
    validate_config(cfg)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


def validate_config(cfg: dict) -> None:
    required_keys = [
        "expected_files",
        "number_regex",
        "left_header_text",
        "total_col_text",
        "total_row_text",
        "mtd_row_text",
        "keep_row_exact",
        "margin_mm",
        "cats",
        "font_name",
        "font_size_pt",
    ]
    for k in required_keys:
        if k not in cfg:
            raise ValueError(f"config.json: відсутній ключ '{k}'")

    if not isinstance(cfg["expected_files"], int) or cfg["expected_files"] <= 0:
        raise ValueError("config.json: expected_files має бути додатнім int")

    if not isinstance(cfg["margin_mm"], int) or cfg["margin_mm"] <= 0:
        raise ValueError("config.json: margin_mm має бути додатнім int")

    if not isinstance(cfg["font_name"], str) or not cfg["font_name"].strip():
        raise ValueError("config.json: font_name має бути непорожнім рядком")

    if not isinstance(cfg["font_size_pt"], int) or cfg["font_size_pt"] <= 0:
        raise ValueError("config.json: font_size_pt має бути додатнім int")

    # regex
    try:
        re.compile(cfg["number_regex"])
    except re.error as e:
        raise ValueError(f"config.json: number_regex некоректний: {e}")

    # cats
    cats = cfg["cats"]
    if not isinstance(cats, list) or len(cats) == 0:
        raise ValueError("config.json: cats має бути списком")
    ids = []
    for item in cats:
        if "id" not in item or "name" not in item:
            raise ValueError("config.json: cats[] має мати 'id' і 'name'")
        ids.append(str(item["id"]).strip())
    if len(ids) != len(set(ids)):
        raise ValueError("config.json: cats[].id мають бути унікальні")


# -------------------------
# Helpers
# -------------------------

def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\u00a0", " ").strip()).lower()


def get_cell_text(cell) -> str:
    return (cell.text or "").strip()


def extract_number(text: str, number_regex: str) -> float:
    m = re.search(number_regex, (text or "").replace("\u00a0", " "))
    if not m:
        return 0.0
    val = m.group(0).replace(",", ".")
    try:
        return float(val)
    except ValueError:
        return 0.0


def fmt_num(x: float) -> str:
    if abs(x - round(x)) < 1e-9:
        return str(int(round(x)))
    return f"{x:.2f}"


def apply_default_font(doc: Document, font_name: str, font_size_pt: int):
    """Шрифт/розмір за замовчуванням + сумісність з кирилицею."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)

    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _apply_run_font(run, cfg: dict, bold: bool):
    run.bold = bold
    run.font.name = cfg["font_name"]
    run.font.size = Pt(int(cfg["font_size_pt"]))
    try:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), cfg["font_name"])
        rFonts.set(qn("w:hAnsi"), cfg["font_name"])
        rFonts.set(qn("w:eastAsia"), cfg["font_name"])
        rFonts.set(qn("w:cs"), cfg["font_name"])
    except Exception:
        pass


def set_cell_center(cell, cfg: dict, bold: bool = False, size_pt: Optional[int] = None, wrap: bool = True):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.word_wrap = bool(wrap)
    font_size = int(size_pt) if size_pt is not None else int(cfg["font_size_pt"])
    local_cfg = {**cfg, "font_size_pt": font_size}

    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            _apply_run_font(r, local_cfg, bold)


def set_cell_left(cell, cfg: dict, bold: bool = False, size_pt: Optional[int] = None, wrap: bool = True):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.word_wrap = bool(wrap)
    font_size = int(size_pt) if size_pt is not None else int(cfg["font_size_pt"])
    local_cfg = {**cfg, "font_size_pt": font_size}

    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            _apply_run_font(r, local_cfg, bold)


def apply_page_setup(doc: Document, margin_mm: int):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def safe_basename(path: str) -> str:
    try:
        return os.path.basename(path)
    except Exception:
        return str(path)


def _apply_heading_style(p, cfg: dict):
    """
    Заголовок:
    - 14pt
    - bold
    - без інтервалу після абзацу
    - кирилиця/шрифт примусово
    """
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)  # ✅ без інтервалів після

    # важливо: Word може робити кілька runs
    for r in p.runs:
        r.bold = True
        r.font.name = cfg["font_name"]
        r.font.size = Pt(14)
        try:
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), cfg["font_name"])
            rFonts.set(qn("w:hAnsi"), cfg["font_name"])
            rFonts.set(qn("w:eastAsia"), cfg["font_name"])
            rFonts.set(qn("w:cs"), cfg["font_name"])
        except Exception:
            pass


# -------------------------
# DOC -> DOCX conversion
# -------------------------

def convert_doc_to_docx(doc_path: str, tmp_dir: str) -> str:
    try:
        import win32com.client  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "Для .doc потрібен встановлений Microsoft Word і pywin32.\n"
            f"pywin32 import error: {e}"
        )

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(doc_path, ReadOnly=True)

        out_path = os.path.join(
            tmp_dir,
            os.path.splitext(os.path.basename(doc_path))[0] + "_converted.docx"
        )
        doc.SaveAs(out_path, FileFormat=16)  # wdFormatXMLDocument
        return out_path
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


# -------------------------
# Table detection/parsing
# -------------------------

@dataclass
class TableIndexes:
    col_unit: int
    col_cat: Dict[str, int]
    col_total: int
    row_first_data: int
    row_total: Optional[int]
    row_mtd: Optional[int]


def find_target_table(doc: Document, cats: List[str], total_col_text: str) -> Optional[Tuple[object, TableIndexes]]:
    for t in doc.tables:
        if len(t.rows) < 3:
            continue

        top = t.rows[0].cells
        bot = t.rows[1].cells
        cols = len(top)
        if cols < 3:
            continue

        col_cat: Dict[str, int] = {}
        for j in range(cols):
            raw = get_cell_text(top[j]).strip()
            if raw in cats:
                col_cat[raw] = j

        if len(col_cat) < max(1, int(len(cats) * 0.8)):
            continue

        col_total = None
        for j in range(cols):
            if norm(get_cell_text(bot[j])) == norm(total_col_text):
                col_total = j
                break
        if col_total is None:
            continue

        col_unit = 0

        row_total = None
        row_mtd = None
        for i in range(2, len(t.rows)):
            name = norm(get_cell_text(t.cell(i, col_unit)))
            if name == norm("Всього"):
                row_total = i
            if "всього" in name and "місяц" in name:
                row_mtd = i

        idx = TableIndexes(
            col_unit=col_unit,
            col_cat=col_cat,
            col_total=col_total,
            row_first_data=2,
            row_total=row_total,
            row_mtd=row_mtd,
        )
        return t, idx

    return None


def parse_one_file(path: str, tmp_dir: str, warnings: List[str], cfg: dict) -> Optional[Dict[str, List[float]]]:
    original = path
    ext = os.path.splitext(path)[1].lower()

    if ext == ".doc":
        try:
            path = convert_doc_to_docx(path, tmp_dir)
        except Exception as e:
            warnings.append(f"[SKIP] {safe_basename(original)}: .doc не оброблено ({e})")
            return None

    try:
        doc = Document(path)
    except Exception as e:
        warnings.append(f"[SKIP] {safe_basename(original)}: не відкрився документ ({e})")
        return None

    cats = [c["id"] for c in cfg["cats"]]
    found = find_target_table(doc, cats=cats, total_col_text=cfg["total_col_text"])
    if not found:
        warnings.append(f"[SKIP] {safe_basename(original)}: не знайшов потрібну таблицю.")
        return None

    table, idx = found
    values: Dict[str, List[float]] = {}

    stop_rows = set(r for r in [idx.row_total, idx.row_mtd] if r is not None)

    keep_row_exact = cfg["keep_row_exact"]
    total_row_text = cfg["total_row_text"]
    mtd_row_text = cfg["mtd_row_text"]

    for r in range(idx.row_first_data, len(table.rows)):
        if r in stop_rows:
            continue

        unit = get_cell_text(table.cell(r, idx.col_unit)).strip()
        if not unit:
            continue

        n_unit = norm(unit)
        if n_unit == norm(total_row_text) or n_unit == norm(mtd_row_text):
            continue

        cat_vals: List[float] = []
        for cat in cats:
            c = idx.col_cat.get(cat)
            cat_vals.append(
                extract_number(get_cell_text(table.cell(r, c)), cfg["number_regex"]) if c is not None else 0.0
            )

        is_keep = n_unit == norm(keep_row_exact)
        total_val = extract_number(get_cell_text(table.cell(r, idx.col_total)), cfg["number_regex"])

        # секційні заголовки відкидаємо, крім keep_row_exact
        if (not is_keep) and all(v == 0.0 for v in cat_vals) and total_val == 0.0:
            continue

        values[unit] = cat_vals

    if not values:
        warnings.append(f"[SKIP] {safe_basename(original)}: не знайшов рядків з даними.")
        return None

    return values


def reorder_units(units: List[str], keep_row_exact: str, total_row_text: str, mtd_row_text: str) -> List[str]:
    filtered = [u for u in units if norm(u) not in (norm(total_row_text), norm(mtd_row_text))]

    target = None
    rest = []
    for u in filtered:
        if norm(u) == norm(keep_row_exact):
            target = u
        else:
            rest.append(u)
    if target is not None:
        rest.append(target)
    return rest


def sum_week(files: List[str], cfg: dict) -> Tuple[Dict[str, List[float]], List[str], List[str]]:
    warnings: List[str] = []
    files_sorted = sorted(files, key=lambda p: os.path.getmtime(p) if os.path.exists(p) else 0)

    cats = [c["id"] for c in cfg["cats"]]
    totals_by_unit: Dict[str, List[float]] = {}
    order: List[str] = []

    with tempfile.TemporaryDirectory(prefix="weekly_tmp_") as tmp_dir:
        for p in files_sorted:
            day = parse_one_file(p, tmp_dir, warnings, cfg)
            if not day:
                continue

            if not order:
                order = list(day.keys())

            for unit, vals in day.items():
                if unit not in totals_by_unit:
                    totals_by_unit[unit] = [0.0] * len(cats)
                    if unit not in order:
                        order.append(unit)
                for i in range(len(cats)):
                    totals_by_unit[unit][i] += vals[i]

    order = reorder_units(order, cfg["keep_row_exact"], cfg["total_row_text"], cfg["mtd_row_text"])
    return totals_by_unit, order, warnings


# -------------------------
# Report generation
# -------------------------

def build_weekly_report_doc(output_path: str, totals_by_unit: Dict[str, List[float]], order: List[str], cfg: dict):
    cats = [c["id"] for c in cfg["cats"]]
    cat_names = {c["id"]: c["name"] for c in cfg["cats"]}

    doc = Document()
    apply_page_setup(doc, cfg["margin_mm"])
    apply_default_font(doc, cfg["font_name"], int(cfg["font_size_pt"]))

    # 2 написи над таблицею (14pt, bold, без інтервалів після)
    p1 = doc.add_paragraph("1. Кількість виявлених кіберінцидентів та порушень захисту інформації за тиждень")
    _apply_heading_style(p1, cfg)

    p2 = doc.add_paragraph("1.1. Розподіл інцидентів кібербезпеки за категоріями")
    _apply_heading_style(p2, cfg)

    # ❌ НЕ додаємо порожній абзац — таблиця йде одразу після p2

    cols_total = 1 + len(cats) + 1
    rows_total = 2 + len(order) + 1

    table = doc.add_table(rows=rows_total, cols=cols_total)
    table.style = "Table Grid"

    # AutoFit по вмісту + центрування таблиці
    table.autofit = True
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # left header merged
    cell00 = table.cell(0, 0)
    cell10 = table.cell(1, 0)
    cell00.merge(cell10)
    cell00.text = cfg["left_header_text"]
    set_cell_center(cell00, cfg, bold=False, size_pt=int(cfg["font_size_pt"]), wrap=True)

    # cats headers (2-row)
    for i, cat in enumerate(cats):
        c = 1 + i
        table.cell(0, c).text = cat
        set_cell_center(table.cell(0, c), cfg, bold=False, size_pt=int(cfg["font_size_pt"]), wrap=True)

        table.cell(1, c).text = cat_names.get(cat, "")
        set_cell_center(table.cell(1, c), cfg, bold=False, size_pt=max(8, int(cfg["font_size_pt"]) - 1), wrap=True)

    # total header merged
    col_total = 1 + len(cats)
    ht0 = table.cell(0, col_total)
    ht1 = table.cell(1, col_total)
    ht0.merge(ht1)
    ht0.text = cfg["total_col_text"]
    ht0.word_wrap = False
    set_cell_center(ht0, cfg, bold=True, size_pt=int(cfg["font_size_pt"]), wrap=False)

    # body
    start_r = 2
    for r_i, unit in enumerate(order, start=start_r):
        table.cell(r_i, 0).text = unit
        set_cell_left(table.cell(r_i, 0), cfg, bold=False, size_pt=int(cfg["font_size_pt"]), wrap=True)

        vals = totals_by_unit.get(unit, [0.0] * len(cats))
        row_sum = 0.0

        for i, v in enumerate(vals):
            row_sum += v
            cell = table.cell(r_i, 1 + i)
            cell.text = fmt_num(v) if v != 0 else ""
            set_cell_center(cell, cfg, bold=False, size_pt=int(cfg["font_size_pt"]), wrap=True)

        tcell = table.cell(r_i, col_total)
        tcell.text = fmt_num(row_sum) if row_sum != 0 else ""
        set_cell_center(tcell, cfg, bold=True, size_pt=int(cfg["font_size_pt"]), wrap=True)

    # bottom TOTAL row
    r_total = start_r + len(order)
    table.cell(r_total, 0).text = cfg["total_row_text"]
    set_cell_left(table.cell(r_total, 0), cfg, bold=True, size_pt=int(cfg["font_size_pt"]), wrap=True)

    col_sums = [0.0] * len(cats)
    for unit in order:
        vals = totals_by_unit.get(unit, [0.0] * len(cats))
        for i in range(len(cats)):
            col_sums[i] += vals[i]

    grand_total = 0.0
    for i, s in enumerate(col_sums):
        grand_total += s
        cell = table.cell(r_total, 1 + i)
        cell.text = fmt_num(s) if s != 0 else ""
        set_cell_center(cell, cfg, bold=True, size_pt=int(cfg["font_size_pt"]), wrap=True)

    gcell = table.cell(r_total, col_total)
    gcell.text = fmt_num(grand_total) if grand_total != 0 else ""
    set_cell_center(gcell, cfg, bold=True, size_pt=int(cfg["font_size_pt"]), wrap=True)

    # "Всього" (останній рядок) — висота як 2 звичайні рядки (AT_LEAST)
    try:
        row = table.rows[r_total]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Mm(12)  # підбирається; якщо мало — Mm(14)
    except Exception:
        pass

    doc.save(output_path)


# -------------------------
# UI
# -------------------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Weekly Report Builder")
        self.geometry("1040x660")
        self.minsize(900, 300)

        self.config_path = DEFAULT_CONFIG_PATH
        self.cfg = self._load_cfg_with_alert()

        self.selected_files: List[str] = []

        # busy state
        self._busy = False
        self._busy_anim_job = None
        self._dot_job = None
        self._dot_phase = 0

        # enterprise palette
        self._C_BG = "#F4F6F9"
        self._C_PANEL = "#FFFFFF"
        self._C_BORDER = "#D0D5DD"
        self._C_TEXT = "#1F2937"
        self._C_MUTED = "#6B7280"
        self._C_ACCENT = "#1D4ED8"
        self._C_ACCENT_HOVER = "#2563EB"
        self._C_OK = "#15803D"

        self.configure(bg=self._C_BG)

        self._apply_theme()
        self._build_header()

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=14, pady=(0, 14))

        self.tab_report = ttk.Frame(nb)
        self.tab_settings = ttk.Frame(nb)

        nb.add(self.tab_report, text="Звіт")
        nb.add(self.tab_settings, text="Налаштування")

        self._build_report_tab()
        self._build_settings_tab()

        # завантаження конфіга у вкладку — БЕЗ алерта на старті
        self.reload_settings(silent=True)

        # progressbar має бути пустий, поки нічого не почалося
        self._set_progress_idle()

    # --------- styling ---------

    def _apply_theme(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except Exception:
            pass

        self.option_add("*Font", ("Segoe UI", 10))

        style.configure("TFrame", background=self._C_PANEL)
        style.configure("Hero.TFrame", background=self._C_BG)
        style.configure("TLabel", background=self._C_PANEL, foreground=self._C_TEXT)
        style.configure("HeroTitle.TLabel", background=self._C_BG, foreground=self._C_TEXT,
                        font=("Segoe UI", 14, "bold"))
        style.configure("HeroSub.TLabel", background=self._C_BG, foreground=self._C_MUTED,
                        font=("Segoe UI", 10))
        style.configure("Muted.TLabel", background=self._C_PANEL, foreground=self._C_MUTED)

        style.configure("TNotebook", background=self._C_BG, borderwidth=0)
        style.configure("TNotebook.Tab",
                        padding=(14, 10),
                        background=self._C_BG,  # фон неактивних вкладок
                        foreground=self._C_MUTED)

        style.map("TNotebook.Tab",
                  background=[
                      ("selected", self._C_PANEL),  # активна вкладка біла
                      ("active", "#EEF2FF")  # hover
                  ],
                  foreground=[
                      ("selected", self._C_TEXT),
                      ("active", self._C_TEXT)
                  ])

        # Прибрати оцю “рамку” навколо контенту вкладки (актуально для clam)
        style.layout("TNotebook", [
            ("TNotebook.client", {"sticky": "nswe"})
        ])

        # buttons
        style.configure(
            "Primary.TButton",
            padding=(14, 10),
            background=self._C_ACCENT,
            foreground="white",
            borderwidth=0,
            focusthickness=0
        )
        style.map(
            "Primary.TButton",
            background=[("active", self._C_ACCENT_HOVER), ("disabled", "#A5B4FC")],
            foreground=[("disabled", "#FFFFFF")]
        )

        style.configure(
            "Secondary.TButton",
            padding=(14, 10),
            background="#E5E7EB",
            foreground=self._C_TEXT,
            borderwidth=0,
            focusthickness=0
        )
        style.map(
            "Secondary.TButton",
            background=[("active", "#D1D5DB"), ("disabled", "#F3F4F6")],
            foreground=[("disabled", self._C_MUTED)]
        )

        # progressbar
        style.configure("TProgressbar", troughcolor="#E5E7EB", background=self._C_ACCENT, borderwidth=0)

    def _build_header(self):
        hero = ttk.Frame(self, style="Hero.TFrame")
        hero.pack(fill="x", padx=14, pady=14)

        left = ttk.Frame(hero, style="Hero.TFrame")
        left.pack(side="left", fill="x", expand=True)


        # стриманий індикатор
        self.dot = tk.Canvas(hero, width=14, height=14, highlightthickness=0, bg=self._C_BG)
        self.dot.pack(side="right", padx=(10, 0))
        self._dot_id = self.dot.create_oval(2, 2, 12, 12, fill=self._C_OK, outline="")

        self._tick_dot()

    def _tick_dot(self):
        self._dot_phase = (self._dot_phase + 1) % 40
        if self._busy:
            r = 5 if self._dot_phase < 20 else 4
            self.dot.coords(self._dot_id, 7 - r, 7 - r, 7 + r, 7 + r)
            self.dot.itemconfig(self._dot_id, fill=self._C_ACCENT)
        else:
            self.dot.coords(self._dot_id, 2, 2, 12, 12)
            self.dot.itemconfig(self._dot_id, fill=self._C_OK)

        self._dot_job = self.after(80, self._tick_dot)

    # --------- core ---------

    def _load_cfg_with_alert(self) -> dict:
        try:
            return load_config(self.config_path)
        except Exception as e:
            messagebox.showerror("Помилка config.json", f"{e}\n\nСтвори/виправ config.json поруч з main.py")
            raise

    def _set_progress_idle(self):
        self.progress.stop()
        self.progress.configure(mode="determinate", value=0, maximum=100)

    def set_busy(self, busy: bool, text: str = ""):
        self._busy = busy

        state = "disabled" if busy else "normal"
        self.btn_pick.config(state=state)
        self.btn_gen.config(state=state)

        if busy:
            self.status.config(text=text)
            self.progress.configure(mode="indeterminate")
            self.progress.start(12)
            self._start_busy_text_anim()
        else:
            self._stop_busy_text_anim()
            self.status.config(text=text)
            self._set_progress_idle()

    def _start_busy_text_anim(self):
        if self._busy_anim_job is not None:
            return

        base = "Обробляю файли"
        self._busy_dots = 0

        def tick():
            if not self._busy:
                self._busy_anim_job = None
                return
            self._busy_dots = (self._busy_dots + 1) % 4
            self.status.config(text=base + ("." * self._busy_dots))
            self._busy_anim_job = self.after(350, tick)

        self._busy_anim_job = self.after(0, tick)

    def _stop_busy_text_anim(self):
        if self._busy_anim_job is not None:
            try:
                self.after_cancel(self._busy_anim_job)
            except Exception:
                pass
            self._busy_anim_job = None

    def _finish_ok(self, save_path: str, skipped: List[str]):
        self.set_busy(False, "Готово")
        if skipped:
            messagebox.showwarning(
                "Готово (частково)",
                f"Звіт збережено:\n{save_path}\n\nПропущено файлів: {len(skipped)}"
            )
        else:
            messagebox.showinfo("Готово", f"Звіт збережено:\n{save_path}")

    def _finish_err(self, msg: str):
        self.set_busy(False, "Помилка")
        messagebox.showerror("Помилка", msg)

    # ---------- Report tab ----------

    def _build_report_tab(self):
        root = ttk.Frame(self.tab_report)
        root.pack(fill="both", expand=True, padx=14, pady=14)

        top = ttk.Frame(root)
        top.pack(fill="x", padx=12, pady=12)

        self.btn_pick = ttk.Button(top, text="Обрати щоденні файли", style="Secondary.TButton", command=self.pick_files)
        self.btn_pick.pack(side="left")

        self.btn_gen = ttk.Button(top, text="Згенерувати звіт", style="Primary.TButton", command=self.generate)
        self.btn_gen.pack(side="left", padx=10)

        self.lbl_info = ttk.Label(root, text="Файли не обрані.", style="Muted.TLabel")
        self.lbl_info.pack(fill="x", padx=12, pady=(0, 6))

        self.progress = ttk.Progressbar(root, mode="determinate", maximum=100, value=0)
        self.progress.pack(fill="x", padx=12, pady=(0, 8))

        self.status = ttk.Label(root, text="", style="Muted.TLabel")
        self.status.pack(fill="x", padx=12, pady=(0, 8))

        self.listbox = tk.Listbox(
            root,
            height=18,
            bg=self._C_PANEL,
            fg=self._C_TEXT,
            highlightthickness=1,
            highlightbackground=self._C_BORDER,
            selectbackground="#DBEAFE",
            selectforeground=self._C_ACCENT,
            activestyle="none",
            relief="solid",
            borderwidth=1
        )
        self.listbox.pack(fill="both", expand=True, padx=12, pady=(0, 10))

        self.lbl_hint = ttk.Label(root, text="", style="Muted.TLabel")
        self.lbl_hint.pack(fill="x", padx=12, pady=(0, 12))

    def pick_files(self):
        files = filedialog.askopenfilenames(
            title="Обери щоденні Word файли",
            filetypes=[
                ("Word documents", "*.docx *.doc"),
                ("DOCX", "*.docx"),
                ("DOC", "*.doc"),
            ],
        )
        if not files:
            return

        self.selected_files = list(files)
        self.listbox.delete(0, tk.END)
        for f in self.selected_files:
            self.listbox.insert(tk.END, f)

        n = len(self.selected_files)
        self.lbl_info.config(text=f"Обрано файлів: {n}")

        if n != int(self.cfg["expected_files"]):
            messagebox.showwarning(
                "Нестандартна кількість файлів",
                f"Обрано {n} файлів, очікувалось {self.cfg['expected_files']}.\n"
                "Можна продовжувати — програма не зламається."
            )

    def generate(self):
        if not self.selected_files:
            messagebox.showerror("Помилка", "Спочатку обери файли .docx/.doc")
            return

        # ЧАС ТІЛЬКИ З СИСТЕМИ (офлайн)
        default_name = datetime.now().strftime("weekly_report_%Y-%m-%d_%H-%M.docx")

        save_path = filedialog.asksaveasfilename(
            title="Куди зберегти тижневий звіт",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile=default_name,
        )
        if not save_path:
            return

        self.set_busy(True, "Обробляю файли")

        def worker():
            try:
                totals_by_unit, order, warnings = sum_week(self.selected_files, self.cfg)
                if not totals_by_unit:
                    raise RuntimeError("Не вдалося витягнути дані з жодного файла.")

                build_weekly_report_doc(save_path, totals_by_unit, order, self.cfg)

                skipped = [w for w in warnings if w.startswith("[SKIP]")]
                self.after(0, lambda: self._finish_ok(save_path, skipped))
            except Exception as e:
                self.after(0, lambda: self._finish_err(str(e)))

        threading.Thread(target=worker, daemon=True).start()

    # ---------- Settings tab ----------

    def _build_settings_tab(self):
        root = ttk.Frame(self.tab_settings)
        root.pack(fill="both", expand=True, padx=14, pady=14)

        top = ttk.Frame(root)
        top.pack(fill="x", padx=12, pady=12)

        ttk.Label(top, text="config.json (можна редагувати прямо тут):").pack(side="left")

        btns = ttk.Frame(top)
        btns.pack(side="right")

        ttk.Button(btns, text="Перезавантажити", style="Secondary.TButton",
                   command=lambda: self.reload_settings(silent=False)).pack(side="left", padx=4)
        ttk.Button(btns, text="Перевірити", style="Secondary.TButton",
                   command=self.validate_settings_text).pack(side="left", padx=4)
        ttk.Button(btns, text="Зберегти", style="Primary.TButton",
                   command=self.save_settings_text).pack(side="left", padx=4)

        self.txt_settings = tk.Text(
            root,
            wrap="none",
            height=24,
            undo=True,
            bg=self._C_PANEL,
            fg=self._C_TEXT,
            insertbackground=self._C_TEXT,
            selectbackground="#DBEAFE",
            highlightthickness=1,
            highlightbackground=self._C_BORDER
        )
        self.txt_settings.pack(fill="both", expand=True, padx=12, pady=(0, 10))

        yscroll = ttk.Scrollbar(root, orient="vertical", command=self.txt_settings.yview)
        yscroll.place(relx=1.0, rely=0.14, relheight=0.74, anchor="ne")
        self.txt_settings.configure(yscrollcommand=yscroll.set)

        xscroll = ttk.Scrollbar(root, orient="horizontal", command=self.txt_settings.xview)
        xscroll.pack(fill="x", padx=12, pady=(0, 12))
        self.txt_settings.configure(xscrollcommand=xscroll.set)

    def reload_settings(self, silent: bool = False):
        """На старті викликаємо silent=True, щоб НЕ було алерта."""
        try:
            cfg = load_config(self.config_path)
            self.cfg = cfg
            self.txt_settings.delete("1.0", tk.END)
            self.txt_settings.insert("1.0", json.dumps(cfg, ensure_ascii=False, indent=2))
            if not silent:
                messagebox.showinfo("Ок", "Налаштування перезавантажено.")
        except Exception as e:
            messagebox.showerror("Помилка", str(e))

    def validate_settings_text(self):
        try:
            raw = self.txt_settings.get("1.0", tk.END).strip()
            cfg = json.loads(raw)
            validate_config(cfg)
            messagebox.showinfo("Ок", "JSON валідний ✅")
        except Exception as e:
            messagebox.showerror("Помилка в JSON", str(e))

    def save_settings_text(self):
        try:
            raw = self.txt_settings.get("1.0", tk.END).strip()
            cfg = json.loads(raw)
            validate_config(cfg)
            save_config(self.config_path, cfg)
            self.cfg = cfg
            messagebox.showinfo("Ок", "Збережено в config.json ✅")
        except Exception as e:
            messagebox.showerror("Помилка збереження", str(e))


if __name__ == "__main__":
    App().mainloop()
